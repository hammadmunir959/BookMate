import os
import shutil
import logging
from typing import List, Dict, Any
from dotenv import load_dotenv
import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for DOCX processing
from pathlib import Path

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_openai import ChatOpenAI
from langchain.chains.question_answering import load_qa_chain

# Load environment variables from .env file
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)

class BookmateRAG:
    ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}

    def __init__(self, 
                 model_name: str = "meta-llama/llama-3-70b-instruct", 
                 embedding_model: str = "sentence-transformers/all-mpnet-base-v2",
                 chunk_size: int = 500, 
                 chunk_overlap: int = 50,
                 uploads_folder: str = "uploads",
                 embeddings_folder: str = "embeddings"):
        """
        Initialize the Bookmate RAG system.

        Args:
            model_name (str): The language model to use.
            embedding_model (str): The embedding model for vectorizing text.
            chunk_size (int): Maximum size of each text chunk.
            chunk_overlap (int): Number of overlapping characters between chunks.
            uploads_folder (str): Folder to store uploaded files.
            embeddings_folder (str): Folder to persist embeddings for each document.
        """
        self.logger = logging.getLogger(self.__class__.__name__)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.uploads_folder = uploads_folder
        self.embeddings_folder = embeddings_folder
        
        # Create necessary directories
        os.makedirs(self.uploads_folder, exist_ok=True)
        os.makedirs(self.embeddings_folder, exist_ok=True)

        # Initialize components
        self.embeddings = HuggingFaceEmbeddings(model_name=embedding_model)
        self.llm = self._load_llm(model_name)
        self.processed_documents: Dict[str, Dict[str, Any]] = {}
        self.current_document: str = None
        
        # Load any existing documents
        self.load_existing_documents()

    def _load_llm(self, model_name: str) -> ChatOpenAI:
        """
        Load the language model from OpenRouter.

        Args:
            model_name (str): The name of the model to load.

        Returns:
            ChatOpenAI: The configured language model instance.
        """
        openrouter_api_key = os.getenv("OPENROUTER_API_KEY")
        if not openrouter_api_key:
            raise ValueError("OpenRouter API key is missing. Ensure .env is properly loaded.")
        
        return ChatOpenAI(
            model_name=model_name,
            openai_api_base="https://openrouter.ai/api/v1",
            openai_api_key=openrouter_api_key,
            temperature=0.7
        )

    def validate_document(self, file_path: str):
        """
        Validate a document to ensure it exists, is not empty, and has an allowed extension.

        Args:
            file_path (str): The path to the document.

        Raises:
            ValueError: If the file does not exist, is empty, or has an invalid extension.
        """
        if not os.path.exists(file_path):
            raise ValueError("File does not exist")
        if os.path.getsize(file_path) == 0:
            raise ValueError("File is empty")
            
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.ALLOWED_EXTENSIONS:
            raise ValueError(f"Invalid file type. Allowed types are: {', '.join(self.ALLOWED_EXTENSIONS)}")

    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from various file types.

        Args:
            file_path (str): The path to the file.

        Returns:
            str: The extracted text content.
        """
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_ext == '.txt':
                return self._extract_text_from_txt(file_path)
            elif file_ext in ['.doc', '.docx']:
                return self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using PyMuPDF.

        Args:
            file_path (str): Path to the PDF file.

        Returns:
            str: Extracted text content.
        """
        doc = fitz.open(file_path)
        text = ""
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text += page.get_text()
        return text

    def _extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file.

        Args:
            file_path (str): Path to the text file.

        Returns:
            str: File contents.
        """
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def _extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX using python-docx.

        Args:
            file_path (str): Path to the Word document.

        Returns:
            str: Extracted text content.
        """
        doc = docx.Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    def process_document(self, file_path: str):
        """
        Process a document: extract text, split into chunks, compute embeddings.

        Args:
            file_path (str): Path to the document to process.
        """
        if file_path in self.processed_documents:
            self.logger.info(f"Document '{file_path}' already processed. Using cached embeddings.")
            return

        self.validate_document(file_path)

        # Extract text using the appropriate method
        doc_text = self.extract_text_from_file(file_path)
        if not doc_text.strip():
            raise ValueError(f"No text extracted from {file_path}. The file might be empty or invalid.")

        # Create document chunks
        documents = [Document(page_content=doc_text)]
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            is_separator_regex=False
        )
        split_documents = text_splitter.split_documents(documents)
        
        # Create and store embeddings
        embeddings_path = os.path.join(self.embeddings_folder, os.path.basename(file_path))
        vector_store = Chroma.from_documents(
            documents=split_documents, 
            embedding=self.embeddings,
            persist_directory=embeddings_path
        )
        
        # Create QA chain
        qa_chain = load_qa_chain(self.llm, chain_type="stuff")
        
        # Cache processed data
        self.processed_documents[file_path] = {
            "split_documents": split_documents,
            "vector_store": vector_store,
            "qa_chain": qa_chain
        }
        self.logger.info(f"Processed document '{file_path}' and cached embeddings.")

    def load_existing_documents(self):
        """
        Load and process all existing files in the uploads folder.
        """
        for file_name in os.listdir(self.uploads_folder):
            file_path = os.path.join(self.uploads_folder, file_name)
            if Path(file_path).suffix.lower() in self.ALLOWED_EXTENSIONS:
                try:
                    self.process_document(file_path)
                    self.logger.info(f"Loaded existing document: {file_path}")
                except Exception as e:
                    self.logger.error(f"Error processing existing document {file_path}: {e}")

    def load_document_embeddings(self, file_path: str) -> Chroma:
        """
        Load the precomputed embeddings for a document.

        Args:
            file_path (str): Path to the document.

        Returns:
            Chroma: The vector store for the document.
        """
        embeddings_path = os.path.join(self.embeddings_folder, os.path.basename(file_path))
        if os.path.exists(embeddings_path):
            vector_store = Chroma(
                persist_directory=embeddings_path,
                embedding_function=self.embeddings
            )
            return vector_store
        else:
            raise ValueError(f"Embeddings for document '{file_path}' not found.")

    def upload_document(self, source_file_path: str) -> str:
        """
        Upload a file to the uploads folder and process it for RAG.

        Args:
            source_file_path (str): The source path of the file.

        Returns:
            str: The destination path where the file was saved.
        """
        file_ext = Path(source_file_path).suffix.lower()
        if file_ext not in self.ALLOWED_EXTENSIONS:
            raise ValueError(f"Invalid file type. Allowed types are: {', '.join(self.ALLOWED_EXTENSIONS)}")

        file_name = os.path.basename(source_file_path)
        dest_path = os.path.join(self.uploads_folder, file_name)

        # Handle duplicate filenames
        counter = 1
        base, ext = os.path.splitext(file_name)
        while os.path.exists(dest_path):
            dest_path = os.path.join(self.uploads_folder, f"{base}_{counter}{ext}")
            counter += 1

        try:
            shutil.copy(source_file_path, dest_path)
            self.logger.info(f"Uploaded file '{file_name}' to '{dest_path}'.")
            self.process_document(dest_path)
            if not self.current_document:
                self.current_document = dest_path
            return dest_path
        except Exception as e:
            if os.path.exists(dest_path):
                os.remove(dest_path)
            self.logger.error(f"Error uploading document {source_file_path}: {e}")
            raise

    def delete_document(self, file_path: str):
        """
        Delete a document and its cached data and embeddings.

        Args:
            file_path (str): Path to the document to delete.
        """
        try:
            # Delete the uploaded file
            if os.path.exists(file_path):
                os.remove(file_path)
                self.logger.info(f"Deleted file '{file_path}'.")
            else:
                self.logger.warning(f"File '{file_path}' not found on disk.")

            # Delete the embeddings folder
            embeddings_path = os.path.join(self.embeddings_folder, os.path.basename(file_path))
            if os.path.exists(embeddings_path):
                shutil.rmtree(embeddings_path)
                self.logger.info(f"Deleted embeddings for '{file_path}'.")
            
            # Remove from cache
            if file_path in self.processed_documents:
                del self.processed_documents[file_path]
                self.logger.info(f"Removed cached data for '{file_path}'.")

            # Reset current document if it was the one deleted
            if self.current_document == file_path:
                self.current_document = None

        except Exception as e:
            self.logger.error(f"Error deleting document {file_path}: {e}")
            raise

    def set_current_document(self, document_path: str):
        """
        Set the current document to be used as context for queries.

        Args:
            document_path (str): Path to the document.
        """
        if document_path not in self.processed_documents:
            # Load embeddings if not already processed
            vector_store = self.load_document_embeddings(document_path)
            self.processed_documents[document_path] = {
                "vector_store": vector_store,
                "qa_chain": load_qa_chain(self.llm, chain_type="stuff")
            }
        self.current_document = document_path
        self.logger.info(f"Current document set to '{document_path}'.")

    def list_documents(self) -> List[str]:
        """
        List all available documents in the uploads folder.

        Returns:
            List[str]: List of file paths.
        """
        return [
            os.path.join(self.uploads_folder, f)
            for f in os.listdir(self.uploads_folder)
            if Path(f).suffix.lower() in self.ALLOWED_EXTENSIONS
        ]

    def query_document(self, query: str) -> str:
        """
        Query the currently selected document using RAG.

        Args:
            query (str): The user's query string.

        Returns:
            str: The generated answer.
        """
        if not self.current_document:
            # No document selected; pass the query directly to the language model
            prompt = f"You are a helpful assistant. Answer the following question: {query}"
            response = self.llm.invoke(prompt)
            return response.content
        
        # Retrieve processed data for the current document
        doc_data = self.processed_documents[self.current_document]
        
        # Use the vector store retriever to find relevant chunks
        retriever = doc_data["vector_store"].as_retriever(
            search_kwargs={"k": min(5, len(doc_data.get('split_documents', [])))}
        )
        relevant_docs = retriever.invoke(query)
        
        # Use the QA chain with the retrieved document chunks
        response = doc_data["qa_chain"].invoke({
            "input_documents": relevant_docs,
            "question": query
        })
        return response.get('output_text', '')

if __name__ == "__main__":
    # Example usage
    rag = BookmateRAG()
    print(f"Initialized BookmateRAG with allowed extensions: {rag.ALLOWED_EXTENSIONS}")